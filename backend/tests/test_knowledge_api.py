"""
知识库接口测试
POST   /api/kb/documents        (文件上传,含类型/大小校验)
GET    /api/kb/documents
DELETE /api/kb/documents/{id}
注:后台文档处理(process_document)已在 conftest 中 mock,避免触发真实 Embedding/Chroma。
"""
from tests.helpers import register_admin_and_login, auth_headers


def _login(client):
    # 知识库管理接口要求 admin 角色,故用 admin 账号登录。
    return register_admin_and_login(client, email="kb@example.com")


def test_upload_no_auth(client):
    files = {"file": ("a.txt", b"hello", "text/plain")}
    r = client.post("/api/kb/documents", files=files)
    assert r.status_code == 401


def test_upload_success_txt(client):
    creds = _login(client)
    h = auth_headers(creds["token"])
    files = {"file": ("doc.txt", b"hello world", "text/plain")}
    r = client.post("/api/kb/documents", files=files, headers=h)
    assert r.status_code == 200
    body = r.json()
    assert body["success"] is True
    data = body["data"]
    # 多文件上传后,data 为列表(单文件也是长度为 1 的列表)
    assert isinstance(data, list) and len(data) == 1
    assert isinstance(data[0]["document_id"], int)
    assert data[0]["status"] == "processing"


def test_upload_success_md(client):
    creds = _login(client)
    h = auth_headers(creds["token"])
    files = {"file": ("doc.md", b"# title\ncontent", "text/markdown")}
    r = client.post("/api/kb/documents", files=files, headers=h)
    assert r.status_code == 200


def test_upload_success_docx(client):
    """Word 文档需被接受并落库为 file_type=docx(依赖 ENUM 已含 docx)。"""
    import io

    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("退换货政策")
    doc.add_paragraph("签收后 7 天内可申请无理由退货。")
    buf = io.BytesIO()
    doc.save(buf)

    creds = _login(client)
    h = auth_headers(creds["token"])
    files = {
        "file": (
            "policy.docx",
            buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    r = client.post("/api/kb/documents", files=files, headers=h)
    assert r.status_code == 200
    assert r.json()["data"][0]["status"] == "processing"

    listed = client.get("/api/kb/documents", headers=h).json()["data"]
    assert any(d["name"] == "policy.docx" and d["file_type"] == "docx" for d in listed)


def test_upload_multiple_files(client):
    """一次请求上传多个文件,全部应落库并在列表中可见。"""
    creds = _login(client)
    h = auth_headers(creds["token"])
    files = [
        ("file", ("a.txt", b"hello", "text/plain")),
        ("file", ("b.md", "# FAQ\nQ1: how to refund?\nA1: apply in 7 days".encode("utf-8"), "text/markdown")),
    ]
    r = client.post("/api/kb/documents", files=files, headers=h)
    assert r.status_code == 200
    body = r.json()
    assert body["success"] is True
    data = body["data"]
    assert isinstance(data, list) and len(data) == 2
    assert all(d["status"] == "processing" for d in data)

    ids = [d["document_id"] for d in data]
    listed = client.get("/api/kb/documents", headers=h).json()["data"]
    assert all(i in [d["id"] for d in listed] for i in ids)


def test_upload_invalid_type(client):
    creds = _login(client)
    h = auth_headers(creds["token"])
    files = {"file": ("malware.exe", b"binary", "application/octet-stream")}
    r = client.post("/api/kb/documents", files=files, headers=h)
    assert r.status_code == 400
    assert r.json()["detail"]["code"] == "INVALID_FILE_TYPE"
    assert "docx" in r.json()["detail"]["message"]


def test_upload_oversize(client):
    creds = _login(client)
    h = auth_headers(creds["token"])
    # 等价类:略大于 10MB 上限
    big = b"x" * (10 * 1024 * 1024 + 1)
    files = {"file": ("big.txt", big, "text/plain")}
    r = client.post("/api/kb/documents", files=files, headers=h)
    assert r.status_code == 400
    assert r.json()["detail"]["code"] == "FILE_TOO_LARGE"


def test_get_documents(client):
    creds = _login(client)
    h = auth_headers(creds["token"])
    up = client.post(
        "/api/kb/documents", files={"file": ("doc.txt", b"hello", "text/plain")}, headers=h
    ).json()["data"]
    r = client.get("/api/kb/documents", headers=h)
    assert r.status_code == 200
    ids = [d["id"] for d in r.json()["data"]]
    assert up[0]["document_id"] in ids


def test_get_documents_no_auth(client):
    r = client.get("/api/kb/documents")
    assert r.status_code == 401


def test_delete_document(client):
    creds = _login(client)
    h = auth_headers(creds["token"])
    up = client.post(
        "/api/kb/documents",
        files={"file": ("doc.txt", b"hi", "text/plain")},
        headers=h,
    ).json()["data"]
    r = client.delete(f"/api/kb/documents/{up[0]['document_id']}", headers=h)
    assert r.status_code == 200
    lst = client.get("/api/kb/documents", headers=h).json()["data"]
    assert all(d["id"] != up["document_id"] for d in lst)


def test_delete_document_not_found(client):
    creds = _login(client)
    h = auth_headers(creds["token"])
    r = client.delete("/api/kb/documents/999999", headers=h)
    assert r.status_code == 404
    assert r.json()["detail"]["code"] == "NOT_FOUND"
