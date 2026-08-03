"""
文档解析(DocumentLoader)单元测试。

覆盖 txt / md / pdf / docx 四种格式的正常解析与异常分支。
PDF 夹具用原始字节手工构造(不依赖 reportlab 等额外库),docx 夹具用
python-docx 生成(已在 requirements 中)。
"""
import zlib

import pytest
from docx import Document as DocxDocument

from app.rag.loader import (
    ALLOWED_UPLOAD_EXTS,
    SUPPORTED_EXTENSIONS,
    DocumentLoader,
)


def _make_pdf(lines, compress=False):
    """构造一个最小可用的单页 PDF(ASCII 文本),xref 偏移按实际长度计算。"""
    text_ops = "BT /F1 12 Tf 72 720 Td 14 TL\n"
    for line in lines:
        text_ops += f"({line}) Tj T*\n"
    text_ops += "ET"
    stream = text_ops.encode("latin-1")

    extra = ""
    if compress:
        stream = zlib.compress(stream)
        extra = " /Filter /FlateDecode"

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        f"<< /Length {len(stream)}{extra} >>\nstream\n".encode("latin-1")
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode("latin-1") + body + b"\nendobj\n"

    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode("latin-1")
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n"
    ).encode("latin-1")
    return bytes(out)


# --------------------------------------------------------------- 支持格式声明
def test_supported_extensions_cover_allowed_uploads():
    """对外声明的可上传后缀必须都能被 loader 分发处理。"""
    for ext in ALLOWED_UPLOAD_EXTS:
        assert ext in SUPPORTED_EXTENSIONS


def test_docx_is_supported():
    assert "docx" in ALLOWED_UPLOAD_EXTS
    assert SUPPORTED_EXTENSIONS["docx"] == "docx"


# --------------------------------------------------------------------- 纯文本
def test_load_txt_utf8(tmp_path):
    p = tmp_path / "a.txt"
    p.write_text("退货政策说明\n七天无理由", encoding="utf-8")
    assert "七天无理由" in DocumentLoader.load(str(p))


def test_load_txt_gbk_fallback(tmp_path):
    p = tmp_path / "gbk.txt"
    p.write_bytes("中文编码回退测试".encode("gbk"))
    assert "中文编码回退测试" in DocumentLoader.load(str(p))


def test_load_md_keeps_markdown(tmp_path):
    p = tmp_path / "a.md"
    p.write_text("# 标题\n- 条目", encoding="utf-8")
    text = DocumentLoader.load(str(p))
    assert text.startswith("# 标题")


# ------------------------------------------------------------------------ PDF
def test_load_pdf_extracts_text(tmp_path):
    p = tmp_path / "a.pdf"
    p.write_bytes(_make_pdf(["Refund policy", "Seven days no reason"]))
    text = DocumentLoader.load(str(p))
    assert "Refund policy" in text
    assert "Seven days no reason" in text


def test_load_pdf_handles_compressed_stream(tmp_path):
    """真实 PDF 内容流通常是 Flate 压缩的,必须能正常解压抽取。"""
    p = tmp_path / "z.pdf"
    p.write_bytes(_make_pdf(["Compressed content stream"], compress=True))
    assert "Compressed content stream" in DocumentLoader.load(str(p))


def test_load_pdf_without_text_layer_raises(tmp_path):
    """扫描件无文本层时必须显式报错,而不是静默返回空串。"""
    p = tmp_path / "empty.pdf"
    p.write_bytes(_make_pdf([]))
    with pytest.raises(ValueError, match="No extractable text"):
        DocumentLoader.load(str(p))


def test_load_pdf_corrupted_raises(tmp_path):
    p = tmp_path / "bad.pdf"
    p.write_bytes(b"this is definitely not a pdf")
    with pytest.raises(ValueError):
        DocumentLoader.load(str(p))


# ----------------------------------------------------------------------- DOCX
def _write_docx(path, *, paragraphs=(), table=None):
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table:
        t = doc.add_table(rows=0, cols=len(table[0]))
        for row in table:
            cells = t.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = value
    doc.save(path)


def test_load_docx_paragraphs(tmp_path):
    p = tmp_path / "a.docx"
    _write_docx(p, paragraphs=["退换货政策", "签收后 7 天内可申请无理由退货。"])
    text = DocumentLoader.load(str(p))
    assert "退换货政策" in text
    assert "签收后 7 天内可申请无理由退货。" in text


def test_load_docx_extracts_tables(tmp_path):
    """表格是业务文档的常见载体,单元格文本必须一并抽取。"""
    p = tmp_path / "t.docx"
    _write_docx(p, table=[["套餐", "价格"], ["基础版", "2999 元/月"]])
    text = DocumentLoader.load(str(p))
    assert "套餐 | 价格" in text
    assert "基础版 | 2999 元/月" in text


def test_load_docx_preserves_document_order(tmp_path):
    """正文与表格需按文档流顺序输出,顺序错乱会破坏分块语义。"""
    p = tmp_path / "order.docx"
    doc = DocxDocument()
    doc.add_paragraph("第一段正文")
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "表格内容"
    doc.add_paragraph("第二段正文")
    doc.save(p)

    text = DocumentLoader.load(str(p))
    assert text.index("第一段正文") < text.index("表格内容") < text.index("第二段正文")


def test_load_docx_skips_blank_paragraphs(tmp_path):
    p = tmp_path / "blank.docx"
    _write_docx(p, paragraphs=["有效内容", "   ", "", "另一段"])
    lines = [ln for ln in DocumentLoader.load(str(p)).splitlines()]
    assert lines == ["有效内容", "另一段"]


def test_load_legacy_doc_raises(tmp_path):
    """旧版二进制 .doc 无法被 python-docx 解析,应给出明确错误。"""
    p = tmp_path / "legacy.docx"
    p.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1legacy ole2 binary")
    with pytest.raises(ValueError, match="legacy .doc is not supported"):
        DocumentLoader.load(str(p))


# ------------------------------------------------------------------ 分发与异常
def test_unsupported_extension_raises(tmp_path):
    p = tmp_path / "a.exe"
    p.write_bytes(b"binary")
    with pytest.raises(ValueError, match="Unsupported file type"):
        DocumentLoader.load(str(p))


def test_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        DocumentLoader.load(str(tmp_path / "nope.txt"))


def test_explicit_file_type_overrides_extension(tmp_path):
    """入库时按 DB 的 file_type 分发,不依赖磁盘上的随机文件名后缀。"""
    p = tmp_path / "8f3a-uuid-no-ext"
    p.write_text("# 标题", encoding="utf-8")
    assert DocumentLoader.load(str(p), file_type="md") == "# 标题"


def test_get_char_count():
    assert DocumentLoader.get_char_count("中文abc") == 5
